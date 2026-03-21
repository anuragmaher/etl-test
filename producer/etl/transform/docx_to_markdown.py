"""Convert a .docx file (as bytes) to Markdown."""

import io

from docx import Document


def convert(docx_bytes: bytes) -> str:
    """Convert .docx file bytes to a Markdown string."""
    doc = Document(io.BytesIO(docx_bytes))
    parts = []

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()

        # Determine heading level
        prefix = ""
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
                prefix = "#" * level + " "
            except ValueError:
                pass

        # Build text with inline formatting
        text_parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue

            stripped = text.strip()
            if not stripped:
                text_parts.append(text)
                continue

            leading = text[: len(text) - len(text.lstrip())]
            trailing = text[len(text.rstrip()) :]
            formatted = stripped

            if run.bold and run.italic:
                formatted = f"***{formatted}***"
            elif run.bold:
                formatted = f"**{formatted}**"
            elif run.italic:
                formatted = f"*{formatted}*"

            text_parts.append(leading + formatted + trailing)

        line = "".join(text_parts)

        # Handle list paragraphs
        if style_name.startswith("list"):
            # Detect nesting from indentation
            indent_level = 0
            if para.paragraph_format.left_indent:
                # Each level is typically 360000 EMUs (0.5 inch)
                indent_level = int(para.paragraph_format.left_indent / 360000)
            indent = "  " * indent_level

            if "bullet" in style_name or "list bullet" in style_name:
                parts.append(f"{indent}- {line}")
            elif "number" in style_name or "list number" in style_name:
                parts.append(f"{indent}1. {line}")
            else:
                parts.append(f"{indent}- {line}")
            continue

        if prefix:
            parts.append(f"{prefix}{line}")
        elif line:
            parts.append(line)
        else:
            parts.append("")

    # Handle tables
    for table in doc.tables:
        parts.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
            if i == 0:
                parts.append("| " + " | ".join(["---"] * len(cells)) + " |")
        parts.append("")

    result = "\n".join(parts)
    while "\n\n\n" in result:
        result = result.replace("\n\n\n", "\n\n")

    return result.strip() + "\n"
